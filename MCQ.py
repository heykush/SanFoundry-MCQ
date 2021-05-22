import requests
from bs4 import BeautifulSoup
import os
# import pandas as pd
# from python-docx import PendingDeprecationWarning

from docx import Document
# from docx.shared import Inches

from selenium import webdriver
from time import sleep
from selenium.webdriver.chrome.options import Options

options = Options()
options.add_extension(r"D:\Download\Chrome extension\addd.crx")

driver = webdriver.Chrome(
    r"C:\Users\gkush\Downloads\Compressed\chromedriver_win32\chromedriver.exe", options=options)
url = "https://www.sanfoundry.com/python-questions-answers-variable-names/"
driver.maximize_window()
driver.get(url)
sleep(4)

# content = driver.page_source.encode('utf-8').strip()
# soup = BeautifulSoup(content,"html.parser")

while True:

    print("\n ==============   QUESTIONS   ================= \n")
    b = []
    e = []
    Questions = driver.find_element_by_xpath('//*[@class="entry-content"]')
    Head = driver.find_element_by_class_name("entry-title")
    Heads = (Head.text)
    Q = (Questions.text)
    C = Q.replace('''Sanfoundry Global Education & Learning Series – Python.
	To practice all areas of Python, here is complete set of 1000+ Multiple Choice Questions and Answers.''', ' ')

    Ques = C.replace('View Answer', '')

    print(Ques)
    heading = (Heads[30:])
    print("\n\n")
    print(" \n=================   ANSWERS  ======================\n")

    Answers = driver.find_elements_by_xpath(
        '//div[@class="collapseomatic_content "][not(li)]')
    index = 1
    for a in Answers:
        ans = a.get_attribute('innerHTML')
        ex = a.get_attribute('innerHTML')[0:9]
        x = ex.replace('Answer:', f'{index}. Answer: ')
        e.extend(x.split("\n"))
        d = ans.replace('<br>', ' ')
        f = d.replace('Answer:', f'{index}. Answer: ')
        b.extend(f.split("\n"))
        index += 1

    Ans = ("\n".join(e))
    print(Ans)

    print("\n\n=================   EXPLANATION  ======================\n")
    Ex = ("\n".join(b))
    print(Ex)

    # To print in excel file
    # df= pd.DataFrame({
    # "Questions": Ques,
    # "Answers": Ans,
    # "Explanation":Ex}, index=[0])
    # df.to_excel("py.xlsx",index=False,header=True)

    # To print in Docs file
    Doc = Document()
    Creator = Doc.add_heading("MCQ's BY GAURAV KUSHWAHA\n", 0)
    Creator.add_run().bold = True
    Doc.add_heading("Multiple Questions\n\n", level=1)
    Doc.add_paragraph(Ques)
    Doc.add_heading("Answer\n\n", level=1)
    Doc.add_paragraph(Ans)
    Doc.add_heading("Explanation\n\n", level=1)
    Doc.add_paragraph(Ex)
    # filepath = r'C:\Users\gkush\My script'+'demo.docx'
    Doc.save('demo.docx')
    # f = open('demo.docx', 'rb')
    # document = Document(f)
    # f.close()

    sleep(2)
    driver.find_element_by_xpath('(//*[contains(text(), "Next")])[1]').click()
    # sleep(3)
